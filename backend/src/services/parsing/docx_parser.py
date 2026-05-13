from docx import Document
from typing import List, Dict, Any

class DocxParser:
    """
    Service for extracting text and basic structure from DOCX files.
    """
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return "\n".join(full_text)
        except Exception as e:
            print(f"Error parsing DOCX {file_path}: {e}")
            return ""

    @staticmethod
    def extract_with_structure(file_path: str) -> List[Dict[str, Any]]:
        """
        Extracts text with heading levels preserved.
        """
        elements = []
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    elements.append({
                        "content": para.text,
                        "style": para.style.name,
                        "is_heading": para.style.name.startswith('Heading')
                    })
            return elements
        except Exception as e:
            print(f"Error parsing structured DOCX {file_path}: {e}")
            return []
